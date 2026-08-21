import json
import os
import re
import io
from typing import List, Optional, Set

import httpx
from fastapi import APIRouter, Request
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from agent.tools.base_knowledge.get_base_knowledge import DOC
from agent.tools.tools_def import llm
from agent.tools.copilot.utils.call_llm_test import call_llm_stream, call_llm
from agent.utils.pd_to_walker import generate_random_string
from agent.utils.get_config import config_data
from data_access.session_log import record_session_operation
from data_access.observe_log import log_observe_cycle
from data_access.report_log import record_report_generation
from utils.front_utils import history_to_text

router = APIRouter()


class DocumentInput(BaseModel):
    conversation_history: List[dict]
    session_id: Optional[str] = None


OUTLINE_SYSTEM = """You are a business document outline generator. Based on the conversation history between a user and a data analysis AI assistant, generate a structured outline for a business summary document.

The outline should:
1. Have a clear, business-oriented title that reflects the analytical goal
2. Break the conversation into logical parts (typically 3-8 parts), focusing on business insights and conclusions
3. Each part should have a heading and a brief description of what business content to cover

IMPORTANT:
- This is a business summary document. The outline should focus on business analysis, data insights, and conclusions. Do NOT include sections about code, SQL queries, technical implementation details, or agent execution process.
- CRITICAL: The output must contain NO code whatsoever. No code blocks, no inline code snippets, no SQL, no Python, no YAML, no chart syntax, no programming language constructs of any kind.
- LANGUAGE IS CRITICAL: The entire document MUST be written in the EXACT SAME language as the user's original question. If the user asked in English, write in English. If the user asked in Chinese, write in Chinese. The conversation history and knowledge base below are provided for factual content ONLY — they may contain mixed or different languages. You MUST ignore their language entirely and write exclusively in the user's language. This is not a suggestion — it is a hard requirement. The language of agent outputs, SQL results, knowledge base, or any other context must NEVER leak into the output.
- The final part (if structured as a summary or conclusion) must synthesize the key findings into concise, actionable insights. It must NOT repeat the headings or structure of the earlier parts. It should answer "so what" and "what to do next," not re-list the document sections. The description of the final part must NOT use phrases like "summarize each section" or "recap the key points from each section" — it should describe a synthesis of cross-cutting conclusions.
- IMPORTANT — CHARTS AND IMAGES: The conversation history contains data analysis results with generated charts and images. When planning the outline, consider which charts are available and assign them to the most relevant parts. Each part should reference at least one chart if suitable charts exist in the context. Do NOT create a separate section just for charts — integrate them naturally into the business analysis sections.

Output ONLY a valid JSON object (no markdown, no code blocks):
{
  "title": "string",
  "parts": [
    {"heading": "string", "description": "string"}
  ]
}"""


PART_SYSTEM = """You are a professional business document writer. Based on the conversation history between a user and a data analysis AI assistant, write the content for a specific section of a business summary document.

Rules:
1. Write in markdown format
2. Focus on business insights, data analysis results, trends, patterns, and conclusions
3. You may mention data sources and filtering criteria when relevant to the business context
4. CRITICAL: The output must contain NO code whatsoever. No code blocks, no inline code snippets, no SQL, no Python, no YAML, no chart syntax, no programming language constructs of any kind.
5. Do NOT describe the agent's execution process, tool calls, or workflow steps
6. CHARTS AND IMAGES ARE REQUIRED: The conversation history contains successfully generated charts and images (URLs like tmp_imgs/*.png). You MUST include every valid chart and image from the context that is relevant to this section's topic. Use markdown image syntax: `![description](image_url)`. Reference the actual image URLs from the conversation history — do NOT make up URLs. Place each chart near the text that discusses its findings. Do NOT repeat the same image across multiple sections — if the prompt lists "already used" images, strictly avoid including them. Each chart should appear in exactly one section (the most relevant one).
7. LANGUAGE IS CRITICAL — READ THIS FIRST: Before writing anything, check the "Document Title" below. If it is in English, write this entire section in English. If it is in Chinese, write in Chinese. The conversation history and knowledge base below are provided for factual content ONLY — they may contain mixed or different languages. You MUST ignore their language entirely and write exclusively in the language of the Document Title. This is not a suggestion — it is a hard requirement. The knowledge base (which is a Chinese document about Singapore regulations) must NEVER influence the output language. Every sentence you write must be in the Document Title's language.
8. Keep the content focused on the section topic
9. Be thorough but concise
10. Use proper markdown headings, lists, and tables as needed. Do NOT use any fenced code blocks. Heading depth is limited to `###` (H3) — do NOT use `####`, `#####`, or `######`. Use `###` for sub-headings and plain bold text or bullet points for deeper structure.
11. Do NOT use `1. xxx` numbered lists (e.g., `1. First item`). Use `-` bullet lists instead. If numbering is essential, use `1) xxx` format to avoid Word auto-numbering conflicts across sections.
11. Sub-headings within this section MAY use numbering (e.g., `### 1. xxx`), but the numbering MUST restart from 1 for THIS section only — do NOT continue numbering from previous sections. Each section is independent; its sub-heading numbers are scoped to this section alone.
12. If this section is the final summary or conclusion: do NOT create sub-sections that mirror the earlier section headings. Do NOT structure the summary as a list of per-topic recaps. Instead, synthesize cross-cutting themes into a few concise, actionable recommendations. Answer "so what" and "what to do next." The summary should be shorter than the other sections, not longer.
"""


def _strip_heading(text: str, heading: str) -> str:
    escaped = re.escape(heading)
    pattern = rf'^\s*(?:#{1,6}\s+)?\*{{0,2}}{escaped}\*{{0,2}}\s*\n+'
    return re.sub(pattern, '', text, count=1)


def _extract_image_urls(text: str) -> Set[str]:
    return set(re.findall(r'!\[[^\]]*\]\(([^)]+)\)', text))


def _download_image_bytes(url: str) -> bytes:
    local_match = re.search(r'tmp_imgs/([^\s/]+\.(?:png|jpg|jpeg|gif|bmp|webp))', url, re.IGNORECASE)
    if local_match:
        local_path = os.path.join("tmp_imgs", local_match.group(1))
        if os.path.isfile(local_path):
            with open(local_path, "rb") as f:
                return f.read()
    try:
        resp = httpx.get(url, timeout=15, follow_redirects=True)
        if resp.status_code == 200:
            return resp.content
    except Exception:
        pass
    return b""


def _markdown_to_docx(markdown_text: str, output_path: str):
    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.font.name = 'Calibri'

    lines = markdown_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            p = doc.add_heading(level=min(level, 3))
            _add_inline_runs(p, text)
            i += 1
            continue

        ul_match = re.match(r'^[-*+]\s+(.+)$', stripped)
        if ul_match:
            p = doc.add_paragraph(style='List Bullet')
            _add_inline_runs(p, ul_match.group(1))
            i += 1
            continue

        ol_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if ol_match:
            p = doc.add_paragraph()
            _add_inline_runs(p, f"{ol_match.group(1)}) {ol_match.group(2)}")
            i += 1
            continue

        img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)$', stripped)
        if img_match:
            alt_text = img_match.group(1)
            img_url = img_match.group(2)
            img_bytes = _download_image_bytes(img_url)
            if img_bytes:
                try:
                    img_stream = io.BytesIO(img_bytes)
                    doc.add_picture(img_stream, width=Inches(5.5))
                    last_paragraph = doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph()
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if alt_text:
                        caption = doc.add_paragraph()
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = caption.add_run(alt_text)
                        run.font.size = Pt(9)
                        run.italic = True
                except Exception:
                    p = doc.add_paragraph()
                    run = p.add_run(f'[Image: {alt_text}]')
                    run.italic = True
            else:
                p = doc.add_paragraph()
                run = p.add_run(f'[Image: {alt_text}]')
                run.italic = True
            i += 1
            continue

        if stripped.startswith('|') and '|' in stripped[1:]:
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_line = lines[i].strip()
                if re.match(r'^[\|\s\-:]+$', row_line):
                    i += 1
                    continue
                cells = [c.strip() for c in row_line.split('|')[1:-1]]
                table_rows.append(cells)
                i += 1
            if table_rows:
                table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                table.style = 'Light Grid Accent 1'
                for ri, row in enumerate(table_rows):
                    for ci, cell_text in enumerate(row):
                        cell = table.rows[ri].cells[ci]
                        cell.text = ''
                        p = cell.paragraphs[0]
                        _add_inline_runs(p, cell_text)
            continue

        p = doc.add_paragraph()
        _add_inline_runs(p, stripped)
        i += 1

    doc.save(output_path)


def _add_inline_runs(paragraph, text: str):
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`([^`]+)`|!\[([^\]]*)\]\(([^)]+)\)|\[([^\]]+)\]\(([^)]+)\))')
    last_end = 0
    for m in pattern.finditer(text):
        if m.start() > last_end:
            paragraph.add_run(text[last_end:m.start()])
        if m.group(2):
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif m.group(3):
            run = paragraph.add_run(m.group(3))
            run.italic = True
        elif m.group(4):
            run = paragraph.add_run(m.group(4))
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        elif m.group(5):
            run = paragraph.add_run(f'[Image: {m.group(5)}]')
            run.italic = True
            run.font.size = Pt(9)
        elif m.group(7):
            run = paragraph.add_run(m.group(7))
            run.underline = True
        last_end = m.end()
    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def _parse_outline_json(raw: str) -> dict:
    raw = raw.strip()
    for prefix in ('```json', '```'):
        if raw.startswith(prefix):
            raw = raw[len(prefix):]
    for suffix in ('```',):
        if raw.endswith(suffix):
            raw = raw[:-len(suffix)]
    raw = raw.strip()
    try:
        result = json.loads(raw)
    except json.JSONDecodeError:
        return {"title": "Summary Document", "parts": []}
    if not isinstance(result, dict):
        return {"title": "Summary Document", "parts": []}
    return {
        "title": result.get("title", "Summary Document"),
        "parts": result.get("parts", [])
    }


def _event_stream_generate_document(conversation_history: List[dict], session_id: str, request_json: str = ""):
    context = history_to_text(conversation_history)

    yield f"data: {json.dumps({'phase': 'outline', 'type': 'msg', 'content': 'Generating document outline...'}, ensure_ascii=False)}\n\n"

    outline_prompt = f"""{OUTLINE_SYSTEM}

{DOC}

Conversation History:
{context}"""

    outline_raw = ""
    for chunk in call_llm_stream(outline_prompt, llm):
        outline_raw += chunk
        yield f"data: {json.dumps({'phase': 'outline', 'type': 'chunk', 'content': chunk}, ensure_ascii=False)}\n\n"

    outline = _parse_outline_json(outline_raw)
    yield f"data: {json.dumps({'phase': 'outline', 'type': 'done', 'content': outline_raw, 'outline': outline}, ensure_ascii=False)}\n\n"

    log_observe_cycle(session_id, 0, "generate_document", "outline",
                      prompt=outline_prompt[:10000], response=outline_raw[:10000],
                      token_estimate=len(outline_prompt) // 3)

    parts = outline.get("parts", [])
    title = outline.get("title", "Summary Document")

    if not parts:
        yield f"data: {json.dumps({'phase': 'outline', 'type': 'error', 'content': 'No parts generated in outline'}, ensure_ascii=False)}\n\n"
        return

    document_parts = []
    used_images: Set[str] = set()
    for i, part in enumerate(parts):
        heading = part.get("heading", f"Part {i + 1}")
        description = part.get("description", "")

        yield f"data: {json.dumps({'phase': 'part', 'type': 'msg', 'content': f'Generating part {i + 1}/{len(parts)}: {heading}', 'part_index': i, 'heading': heading}, ensure_ascii=False)}\n\n"

        used_hint = ""
        if used_images:
            used_list = "\n".join(f"- {url}" for url in sorted(used_images))
            used_hint = f"\n\nIMPORTANT: The following images have already been used in previous sections. Do NOT include them again:\n{used_list}\n"

        part_prompt = f"""{PART_SYSTEM}

Document Title: {title}
Section Heading: {heading}
Section Description: {description}{used_hint}

Full Document Outline (all sections):
{chr(10).join(f'  {j+1}. {p["heading"]} — {p["description"]}' for j, p in enumerate(parts))}

{DOC}

Conversation History:
{context}

Write the content for the section "{heading}" in markdown format. ⚠️ CRITICAL: Do NOT repeat the section heading "{heading}" in your output. The heading will be added automatically. Start directly with the body content — no heading, no title line."""

        part_raw = ""
        for chunk in call_llm_stream(part_prompt, llm):
            part_raw += chunk
            yield f"data: {json.dumps({'phase': 'part', 'type': 'chunk', 'content': chunk, 'part_index': i}, ensure_ascii=False)}\n\n"

        part_raw = _strip_heading(part_raw, heading)

        new_images = _extract_image_urls(part_raw)
        used_images.update(new_images)

        document_parts.append((heading, part_raw))
        yield f"data: {json.dumps({'phase': 'part', 'type': 'done', 'content': part_raw, 'part_index': i, 'heading': heading}, ensure_ascii=False)}\n\n"

        log_observe_cycle(session_id, i + 1, "generate_document", "part",
                          prompt=part_prompt[:10000], response=part_raw[:10000],
                          token_estimate=len(part_prompt) // 3)

    full_document = f"# {title}\n\n"
    for idx, (heading, content) in enumerate(document_parts, 1):
        full_document += f"## {idx}. {heading}\n\n{content}\n\n"

    full_document = re.sub(r'```[a-z]*\n.*?```\n?', '', full_document, flags=re.DOTALL)

    file_name = f"doc_{generate_random_string(8)}"
    os.makedirs("tmp_imgs", exist_ok=True)

    md_path = os.path.join("tmp_imgs", file_name + ".md")
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(full_document)

    docx_path = os.path.join("tmp_imgs", file_name + ".docx")
    _markdown_to_docx(full_document, docx_path)

    static_url = config_data.get("static_path", "http://127.0.0.1:8009/")
    download_url_md = static_url.rstrip("/") + "/tmp_imgs/" + file_name + ".md"
    download_url_docx = static_url.rstrip("/") + "/tmp_imgs/" + file_name + ".docx"

    yield f"data: {json.dumps({'phase': 'document', 'type': 'done', 'content': full_document, 'title': title, 'parts_count': len(parts), 'file_name': file_name, 'download_url_md': download_url_md, 'download_url_docx': download_url_docx}, ensure_ascii=False)}\n\n"

    record_session_operation(
        session_id, "/api/generate-document/stream/",
        request_json, full_document[:5000], "",
        "success", f"文档生成完成: {title}, 共{len(parts)}部分",
        prompt_length=len(context)
    )

    record_report_generation(
        session_id=session_id,
        file_name=file_name,
        chat_history=json.dumps(conversation_history, ensure_ascii=False),
        outline=outline_raw,
        full_text=full_document,
    )


@router.post("/api/generate-document/stream/")
async def generate_document_stream_api(request: Request, user_input: DocumentInput):
    return StreamingResponse(
        _event_stream_generate_document(
            user_input.conversation_history,
            user_input.session_id or "",
            user_input.model_dump_json(),
        ),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        }
    )