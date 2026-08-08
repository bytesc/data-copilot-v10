import time
import json
from typing import Optional, List
import io
import base64
import httpx
import random
import string
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pywebio.session import set_env
from pywebio.input import input, TEXT, textarea, file_upload, select, checkbox
from pywebio.output import put_text, put_html, put_markdown, clear, put_loading, toast, popup, put_buttons, \
    put_collapse, put_table, use_scope, put_scope
from pywebio import start_server, config
from data_access.read_db import get_rows_from_all_tables, get_table_comments_dict, get_all_comments_from_table
from utils.get_config import config_data
import markdown
from bs4 import BeautifulSoup
import requests


SELECT_TABLES = []
SELECT_LABELS = []


def ai_agent_api(question: str, tables: Optional[List[str]] = None, path: str = "/api/ask-agent/",
                 url="http://127.0.0.1:" + str(config_data["server_port"]), session_id: str = ""):
    with httpx.Client(timeout=300.0) as client:
        try:
            payload = {"question": question, "session_id": session_id}
            if tables:
                payload["tables"] = tables

            response = client.post(url + path, json=payload)
            if response.status_code == 200:
                print(response.json()["ans"])
                return response.json()["ans"], response.json()["code"], response.json().get("session_id", "")
            else:
                return None, None, session_id
        except httpx.RequestError as e:
            print(e)
            return None, None, session_id


def generate_code_stream(question: str, tables: Optional[List[str]] = None,
                         selected_fields: Optional[dict] = None,
                         url="http://127.0.0.1:" + str(config_data["server_port"]), session_id: str = ""):
    payload = {"question": question, "tables": tables, "session_id": session_id}
    if selected_fields:
        payload["selected_fields"] = selected_fields
    with httpx.stream("POST", url + "/api/generate-code/stream/",
                      json=payload,
                      timeout=300.0) as response:
        if response.status_code != 200:
            yield {"type": "error", "content": f"HTTP {response.status_code}"}
            return

        full_code = ""
        buffer = ""
        for chunk in response.iter_text():
            buffer += chunk
            while "\n\n" in buffer:
                event_str, buffer = buffer.split("\n\n", 1)
                for line in event_str.split("\n"):
                    if line.startswith("data: "):
                        try:
                            event = json.loads(line[6:])
                            if event.get("type") == "code_chunk":
                                full_code += event["content"]
                            if event.get("type") == "code_complete":
                                full_code = event["content"]
                            yield event
                        except json.JSONDecodeError:
                            pass

        yield {"type": "done", "content": "", "full_code": full_code}


def execute_code_stream(code: str,
                        url="http://127.0.0.1:" + str(config_data["server_port"]), session_id: str = ""):
    with httpx.stream("POST", url + "/api/exe-code/stream/",
                      json={"code": code, "session_id": session_id},
                      timeout=300.0) as response:
        if response.status_code != 200:
            yield {"type": "error", "content": f"HTTP {response.status_code}"}
            return

        full_ans = ""
        buffer = ""
        for chunk in response.iter_text():
            buffer += chunk
            while "\n\n" in buffer:
                event_str, buffer = buffer.split("\n\n", 1)
                for line in event_str.split("\n"):
                    if line.startswith("data: "):
                        try:
                            event = json.loads(line[6:])
                            if event.get("type") == "chunk":
                                full_ans += event["content"]
                            yield event
                        except json.JSONDecodeError:
                            pass

        yield {"type": "done", "content": "", "full_ans": full_ans}


def step_chat_api_stream(question: str,
                         url="http://127.0.0.1:" + str(config_data["server_port"]), session_id: str = ""):
    with httpx.stream("POST", url + "/api/step-chat/stream/",
                      json={"question": question, "session_id": session_id},
                      timeout=300.0) as response:
        if response.status_code != 200:
            yield {"type": "error", "content": f"HTTP {response.status_code}"}
            return

        full_ans = ""
        buffer = ""
        for chunk in response.iter_text():
            buffer += chunk
            while "\n\n" in buffer:
                event_str, buffer = buffer.split("\n\n", 1)
                for line in event_str.split("\n"):
                    if line.startswith("data: "):
                        try:
                            event = json.loads(line[6:])
                            if event.get("type") == "chunk":
                                full_ans += event["content"]
                            yield event
                        except json.JSONDecodeError:
                            pass

        yield {"type": "done", "content": "", "full_ans": full_ans}


def filter_db_fields_stream(question: str, tables: Optional[List[str]] = None,
                            url="http://127.0.0.1:" + str(config_data["server_port"])):
    with httpx.stream("POST", url + "/api/filter-db-fields/stream/",
                      json={"question": question, "tables": tables},
                      timeout=300.0) as response:
        if response.status_code != 200:
            yield {"type": "error", "content": f"HTTP {response.status_code}"}
            return

        full_content = ""
        buffer = ""
        for chunk in response.iter_text():
            buffer += chunk
            while "\n\n" in buffer:
                event_str, buffer = buffer.split("\n\n", 1)
                for line in event_str.split("\n"):
                    if line.startswith("data: "):
                        try:
                            event = json.loads(line[6:])
                            if event.get("type") == "chunk":
                                full_content += event["content"]
                            yield event
                        except json.JSONDecodeError:
                            pass

        yield {"type": "done", "content": full_content}


def upload_csv_api(file_content, table_name="uploaded_data"):
    url = f"http://127.0.0.1:{config_data['server_port']}/upload-csv/"
    files = {
        'file': ('data.csv', file_content, 'text/csv')
    }
    data = {
        'table_name': table_name
    }
    with httpx.Client(timeout=60.0) as client:
        try:
            response = client.post(url, files=files, data=data)
            if response.status_code == 200:
                return response.json()
            else:
                return {"error": f"{response.status_code}", "details": response.text}
        except httpx.RequestError as e:
            return {"error": f"{str(e)}"}


def upload_doc_api(file_content, filename, table_name="uploaded_data"):
    url = f"http://127.0.0.1:{config_data['server_port']}/upload-txt/"
    files = {
        'file': (filename, file_content, 'application/octet-stream')
    }
    data = {
        'table_name': table_name
    }
    with httpx.Client(timeout=60.0) as client:
        try:
            response = client.post(url, files=files, data=data)
            if response.status_code == 200:
                return response.json()
            else:
                return {"error": f"{response.status_code}", "details": response.text}
        except httpx.RequestError as e:
            return {"error": f"{str(e)}"}


def download_image(url):
    """下载网络图片"""
    try:
        response = requests.get(url, timeout=30)
        if response.status_code == 200:
            return io.BytesIO(response.content)
    except:
        pass
    return None


def markdown_to_word(doc, markdown_text):
    """将Markdown转换为Word文档内容"""
    html = markdown.markdown(markdown_text, extensions=['extra', 'tables'])
    soup = BeautifulSoup(html, 'html.parser')

    for element in soup.children:
        if element.name == 'h1':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'h4':
            doc.add_heading(element.get_text(), level=4)
        elif element.name == 'p':
            p = doc.add_paragraph()
            for child in element.children:
                if child.name == 'strong' or child.name == 'b':
                    run = p.add_run(child.get_text())
                    run.bold = True
                elif child.name == 'em' or child.name == 'i':
                    run = p.add_run(child.get_text())
                    run.italic = True
                elif child.name == 'a':
                    run = p.add_run(child.get_text())
                    run.underline = True
                elif child.name == 'img':
                    img_url = child.get('src')
                    if img_url:
                        img_data = download_image(img_url)
                        if img_data:
                            try:
                                doc.add_picture(img_data, width=Inches(5))
                            except:
                                p.add_run(f"[图片: {img_url}]")
                        else:
                            p.add_run(f"[图片加载失败: {img_url}]")
                elif child.string:
                    p.add_run(child.string)
                elif child.name is None:
                    if child.strip():
                        p.add_run(child.strip())
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                doc.add_paragraph(li.get_text(), style='List Bullet')
        elif element.name == 'ol':
            for idx, li in enumerate(element.find_all('li', recursive=False), 1):
                doc.add_paragraph(f"{idx}. {li.get_text()}", style='List Number')
        elif element.name == 'table':
            rows = element.find_all('tr')
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['td', 'th'])))
                table.style = 'Table Grid'
                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell.get_text().strip()
        elif element.name == 'blockquote':
            p = doc.add_paragraph()
            p.add_run(element.get_text()).italic = True
            p.paragraph_format.left_indent = Inches(0.5)
        elif element.name == 'hr':
            doc.add_page_break()
        elif element.name == 'br':
            doc.add_paragraph()
        elif element.string and element.string.strip():
            doc.add_paragraph(element.string.strip())


def export_full_to_word(conversation_history):
    """导出完整的对话历史为Word文档"""
    doc = Document()

    title = doc.add_heading('Data-Copilot Conversation Export (Full)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    from datetime import datetime
    doc.add_paragraph(f"Export Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()

    for entry in conversation_history:
        if entry.startswith('Q: '):
            doc.add_heading('Question:', level=2)
            doc.add_paragraph(entry[3:])
        elif entry.startswith('A: '):
            doc.add_heading('Answer:', level=2)
            markdown_to_word(doc, entry[3:])
        elif entry.startswith('Code Generated: '):
            doc.add_heading('Generated Code:', level=2)
            code_para = doc.add_paragraph()
            code_run = code_para.add_run(entry[16:])
            code_run.font.name = 'Courier New'
            code_run.font.size = Pt(10)
        elif entry.startswith('Exe Result: '):
            doc.add_heading('Execution Result:', level=2)
            markdown_to_word(doc, entry[12:])
        elif entry.startswith('Planner: '):
            doc.add_heading('Plan:', level=2)
            markdown_to_word(doc, entry[9:])
        else:
            doc.add_paragraph(entry)

        doc.add_paragraph('_' * 50)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream


def export_essentials_to_word(conversation_history):
    """只导出答案(ans)和第一个问题"""
    doc = Document()

    title = doc.add_heading('Data-Copilot Export', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    from datetime import datetime
    doc.add_paragraph(f"Export Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()

    first_question = None
    answers = []

    for entry in conversation_history:
        if entry.startswith('Q: ') and first_question is None:
            first_question = entry[3:]
        elif entry.startswith('A: '):
            answers.append(entry[3:])

    if first_question:
        doc.add_paragraph(first_question)
        doc.add_paragraph()

    if answers:
        for answer in answers:
            markdown_to_word(doc, answer)
            doc.add_paragraph()

    if not first_question and not answers:
        doc.add_paragraph("No essential content found to export.")

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream


def handle_export_word(conversation_history, export_type="full"):
    """处理导出Word文档"""
    if not conversation_history:
        toast("No content to export!", color='warning')
        return

    with put_loading(shape="grow", color="primary"):
        try:
            if export_type == "full":
                word_file = export_full_to_word(conversation_history)
                filename = "conversation_export_full.docx"
            else:
                word_file = export_essentials_to_word(conversation_history)
                filename = "conversation_export_essentials.docx"

            file_content = word_file.getvalue()
            b64 = base64.b64encode(file_content).decode()

            download_script = f'''
            <script>
                var link = document.createElement('a');
                link.href = 'data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}';
                link.download = '{filename}';
                document.body.appendChild(link);
                link.click();
                document.body.removeChild(link);
            </script>
            '''

            put_html(download_script)
            toast(f"Download started: {filename}", color='success')

        except Exception as e:
            toast(f"Export failed: {str(e)}", color='error')


def handle_csv_upload():
    file_info = file_upload(
        "Please select a CSV file to upload",
        accept=".csv",
        help_text="Select the CSV file you want to upload"
    )

    if file_info:
        table_name = input("Enter table name (optional, default is 'uploaded_data')", type=TEXT,
                           placeholder="uploaded_data", required=False)
        if not table_name:
            table_name = "uploaded_data"
        with put_loading(shape="grow", color="primary"):
            result = upload_csv_api(file_info['content'], table_name)
            print(result)
        err = result.get('type', "error")
        if err == "error":
            toast(f"Upload failed: {result}", color='error')
        else:
            toast("File uploaded successfully!", color='success')
            put_markdown("### Upload Results")
            put_markdown(f"Table name: `{result.get('table_name', table_name)}`")
            put_markdown(f"Row count: {result.get('row_count', 'N/A')}")
            put_markdown(f"Message: {result.get('message', 'N/A')}")


def handle_doc_upload():
    file_info = file_upload(
        "Please select a document file to upload (txt, doc, docx, pdf)",
        accept=".txt,.doc,.docx,.pdf",
        help_text="Select the document file you want to upload"
    )

    if file_info:
        table_name = input("Enter table name (optional, default is 'uploaded_data')", type=TEXT,
                           placeholder="uploaded_data", required=False)
        if not table_name:
            table_name = "uploaded_data"
        with put_loading(shape="grow", color="primary"):
            result = upload_doc_api(file_info['content'], file_info['filename'], table_name)
            print(result)

        if result.get('error'):
            toast(f"Upload failed: {result.get('error')}", color='error')
        else:
            toast("File uploaded successfully!", color='success')
            put_markdown("### Upload Results")
            put_markdown(f"Table name: `{result.get('table_name', table_name)}`")
            put_markdown(f"Extracted text length: {result.get('extracted_text_length', 'N/A')}")
            put_markdown(f"Preview: {result.get('preview', 'N/A')}")


def handle_table_selection(table_options):
    global SELECT_TABLES, SELECT_LABELS
    checkbox_options = [(opt['label'], opt['value']) for opt in table_options]
    selected_tables = checkbox(
        "Select tables: ",
        options=checkbox_options,
        inline=True
    )
    SELECT_TABLES = selected_tables
    put_markdown(f"You have selected: `{', '.join(selected_tables)}`")
    if selected_tables:
        selected_labels = []
        for table_value in selected_tables:
            for opt in table_options:
                if opt['value'] == table_value:
                    selected_labels.append(opt['label'])
                    break
        SELECT_LABELS = selected_labels


def display_streaming_response(scope_name):
    """返回一个回调函数，用于在指定scope中追加显示流式内容"""
    accumulated = ""

    def append_chunk(event):
        nonlocal accumulated
        event_type = event.get("type")
        if event_type == "status":
            toast(event["content"], color='info')
        elif event_type in ("chunk", "code_chunk"):
            accumulated += event["content"]
            with use_scope(scope_name, clear=True):
                put_markdown(accumulated, sanitize=False)
        elif event_type == "code_complete":
            accumulated = "```python\n" + event["content"] + "\n```"
            with use_scope(scope_name, clear=True):
                put_collapse("📝 Generated Code", [put_markdown(accumulated, sanitize=False)])
        elif event_type == "solved":
            accumulated = event["content"]
            with use_scope(scope_name, clear=True):
                put_markdown(accumulated, sanitize=False)
        elif event_type == "error":
            toast(event["content"], color='error')
            with use_scope(scope_name, clear=False):
                put_markdown(f"\n\n**Error:** {event['content']}", sanitize=False)
        return accumulated

    return append_chunk


def main():
    global SELECT_TABLES, SELECT_LABELS
    put_markdown("# Data-Copilot (Streaming)")

    first_five_rows = get_rows_from_all_tables()

    table_comments = get_table_comments_dict()
    table_options = []
    for table_name, comment in table_comments.items():
        display_name = f"{table_name} ({comment})" if comment else table_name
        table_options.append({'label': display_name, 'value': table_name})

    put_markdown("### Control Panel")
    put_buttons(['Select Tables', 'Upload CSV File', 'Upload Document File'],
                onclick=[lambda: handle_table_selection(table_options), handle_csv_upload, handle_doc_upload])

    put_markdown("### Export Options")
    put_buttons(['Export Full Conversation', 'Export Essentials (Answers)'],
                onclick=[lambda: handle_export_word(conversation_history, "full"),
                         lambda: handle_export_word(conversation_history, "essentials")])

    put_markdown("### 📊 Data View")
    with put_collapse(f"📋 Tables"):
        all_comments = get_all_comments_from_table()
        first_five_rows = get_rows_from_all_tables()

        for table_name, rows in first_five_rows.items():
            with put_collapse(f" table {table_name}"):
                if table_name in all_comments:
                    table_comment = all_comments[table_name].get('table_comment', '')
                    if table_comment:
                        put_text(f"📝 {table_comment}")

                    columns = all_comments[table_name].get('columns', {})
                    if columns:
                        comment_table = [["Column Name", "Comment"]]
                        for col_name, comment in columns.items():
                            comment_table.append([col_name, comment])
                        put_table(comment_table)

                put_text(f"📊 table {table_name} first 5 rows:")
                put_table([rows.columns.tolist()] + rows.values.tolist())

    conversation_history = []
    conversation_session_id = datetime.now().strftime("%Y%m%d%H%M%S") + "".join(random.choices(string.ascii_letters, k=8))
    put_markdown(conversation_session_id)

    question = textarea("Enter your question here:", type=TEXT, rows=2)
    put_markdown("## " + question)
    conversation_history.append(f"Q: {question}")

    scope_name = f"plan_scope"
    put_scope(scope_name)
    append_plan = display_streaming_response(scope_name)
    full_plan = ""
    for event in step_chat_api_stream(question, session_id=conversation_session_id):
        full_plan = append_plan(event)
    if full_plan:
        conversation_history.append(f"Planner: {full_plan}")
    else:
        put_text("Failed to get a response from the AI Agent.")

    while True:
        table_pre = ""

        value = "please do the next step on the todo list"
        question = textarea("What is next?:", value=value, type=TEXT, rows=2)
        put_markdown("## " + question)
        if conversation_history:
            context = "\n".join(conversation_history)
            full_question = f"Context:\n{context}\n\nCurrent Question:\n{question}"
        else:
            full_question = question

        if value == question:
            filter_scope = f"filter_scope_{len(conversation_history)}"
            put_scope(filter_scope)
            filter_content = ""
            selected_fields = None
            for event in filter_db_fields_stream(table_pre + full_question, SELECT_TABLES):
                event_type = event.get("type")
                if event_type == "status":
                    toast(event["content"], color='info')
                elif event_type == "chunk":
                    filter_content += event["content"]
                    with use_scope(filter_scope, clear=True):
                        put_collapse("🔍 数据库字段筛选", [put_markdown("```json\n" + filter_content + "\n```", sanitize=False)])
                elif event_type == "done":
                    filter_content = event.get("content", "")
                    import re as _re
                    match = _re.search(r'```json\s*(.*?)\s*```', filter_content, _re.DOTALL)
                    if match:
                        try:
                            selected_fields = json.loads(match.group(1))
                        except json.JSONDecodeError:
                            pass
                    with use_scope(filter_scope, clear=True):
                        if selected_fields:
                            put_collapse("🔍 数据库字段筛选", [
                                put_markdown("✅ 已筛选以下表和字段：", sanitize=False),
                                put_markdown("```json\n" + json.dumps(selected_fields, ensure_ascii=False, indent=2) + "\n```", sanitize=False)
                            ])
                        else:
                            put_collapse("🔍 数据库字段筛选", [put_markdown("```json\n" + filter_content + "\n```", sanitize=False)])
                elif event_type == "error":
                    toast(event["content"], color='warning')

            code_scope = f"code_scope_{len(conversation_history)}"
            put_scope(code_scope)
            append_code = display_streaming_response(code_scope)
            full_code = ""
            solved_ans = ""
            for event in generate_code_stream(table_pre + full_question, SELECT_TABLES,
                                              selected_fields=selected_fields,
                                              session_id=conversation_session_id):
                append_code(event)
                if event.get("type") == "code_complete":
                    full_code = event["content"]
                elif event.get("type") == "solved":
                    solved_ans = event["content"]

            if full_code:
                ans_scope = f"ans_scope_{len(conversation_history)}"
                put_scope(ans_scope)
                append_ans = display_streaming_response(ans_scope)
                full_ans = ""
                for event in execute_code_stream(full_code, session_id=conversation_session_id):
                    full_ans = append_ans(event)
                if full_ans:
                    conversation_history.append(f"Q: {question}")
                    conversation_history.append(f"A: {full_ans}")
            elif solved_ans:
                conversation_history.append(f"Q: {question}")
                conversation_history.append(f"A: {solved_ans}")
            else:
                put_text("Failed to get a response from the AI Agent.")

            context = "\n".join(conversation_history)
            full_question = f"Context:\n{context}\n"
        else:
            context = "\n".join(conversation_history)
            full_question = f"Context:\n{context}\n\nCurrent Question:\n{question}"

        plan_scope = f"plan_scope_{len(conversation_history)}"
        put_scope(plan_scope)
        append_plan = display_streaming_response(plan_scope)
        full_plan = ""
        for event in step_chat_api_stream(table_pre + full_question, session_id=conversation_session_id):
            full_plan = append_plan(event)
        if full_plan:
            conversation_history.append(f"Planner: {full_plan}")
        else:
            put_text("Failed to get a response from the AI Agent.")


if __name__ == '__main__':
    start_server(main, port=8038, debug=True)
