import time
import json
from typing import Optional, List, Dict
import io
import base64
import httpx
import random
import string
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
from bs4 import BeautifulSoup
import requests
from utils.get_config import config_data


FRONTEND_ACTIONS = {"output_text", "ask_question", "ask_choice", "summary_and_pause", "attempt_completion"}


def history_to_text(history: List[dict]) -> str:
    lines = []
    for entry in history:
        role = entry.get("role", "")
        entry_type = entry.get("type", "")
        if role == "user":
            utype = entry.get("type", "question")
            content = entry.get("content", "")
            if utype == "question":
                lines.append(f"Q: {content}")
            elif utype == "choice":
                lines.append(f"User chose: {content}")
            elif utype == "response":
                lines.append(f"User response: {content}")
            elif utype == "input":
                lines.append(f"User: {content}")
        elif entry_type == "think":
            lines.append(f"[THINK] Plan:\n{entry.get('content', '')}")
        elif entry_type == "action_decision":
            lines.append(f"[ACTION] Decision:\n{entry.get('content', '')}")
        elif entry_type == "act":
            action = entry.get("action", "")
            if action == "explore_schema":
                if entry.get("selected_fields") is not None:
                    lines.append(f"[ACT explore_schema] Selected Fields: {json.dumps(entry['selected_fields'], ensure_ascii=False)}")
                if entry.get("explore_plan"):
                    lines.append(f"[ACT explore_schema] Query Plan:\n{entry['explore_plan']}")
                if entry.get("search_result"):
                    lines.append(f"[ACT explore_schema] Results:\n{entry['search_result']}")
            elif action == "explore_functions":
                if entry.get("selected_functions") is not None:
                    lines.append(f"[ACT explore_functions] Selected Functions: {json.dumps(entry['selected_functions'], ensure_ascii=False)}")
                if entry.get("search_result"):
                    lines.append(f"[ACT explore_functions] Results:\n{entry['search_result']}")
            elif action == "generate_and_execute":
                if entry.get("code"):
                    lines.append(f"[ACT generate_and_execute] Code:\n{entry['code']}")
                if entry.get("error"):
                    lines.append(f"[ACT generate_and_execute] Error:\n{entry['error']}")
                elif entry.get("result"):
                    lines.append(f"[ACT generate_and_execute] Result:\n{entry['result']}")
            elif action == "solved":
                if entry.get("solved_ans"):
                    lines.append(f"[ACT] Solved Answer:\n{entry['solved_ans']}")
            elif action in FRONTEND_ACTIONS:
                if entry.get("text"):
                    lines.append(f"[ACT {action}] Output:\n{entry['text']}")
            else:
                if entry.get("search_result"):
                    lines.append(f"[ACT {action}] Results:\n{entry['search_result']}")
        elif entry_type == "observe":
            action = entry.get("action", "")
            lines.append(f"[OBSERVE {action}] Review:\n{entry.get('content', '')}")
    return "\n".join(lines)


def ai_agent_api(question: str, tables: Optional[List[str]] = None, path: str = "/api/ask-agent/",
                 url="http://127.0.0.1:" + str(config_data["server_port"]), session_id: str = ""):
    with httpx.Client(timeout=300.0) as client:
        try:
            payload = {"question": question, "session_id": session_id}
            if tables:
                payload["tables"] = tables

            response = client.post(url + path, json=payload)
            if response.status_code == 200:
                print(response.json()["ans"])
                return response.json()["ans"], response.json()["code"], response.json().get("session_id", "")
            else:
                return None, None, session_id
        except httpx.RequestError as e:
            print(e)
            return None, None, session_id


def generate_code_stream(question: str, tables: Optional[List[str]] = None,
                         selected_fields: Optional[dict] = None,
                         selected_functions: Optional[List[str]] = None,
                         url="http://127.0.0.1:" + str(config_data["server_port"]), session_id: str = ""):
    payload = {"question": question, "tables": tables, "session_id": session_id}
    if selected_fields:
        payload["selected_fields"] = selected_fields
    if selected_functions:
        payload["selected_functions"] = selected_functions
    with httpx.stream("POST", url + "/api/generate-code/stream/",
                      json=payload,
                      timeout=300.0) as response:
        if response.status_code != 200:
            yield {"type": "error", "content": f"HTTP {response.status_code}"}
            return

        full_code = ""
        buffer = ""
        for chunk in response.iter_text():
            buffer += chunk
            while "\n\n" in buffer:
                event_str, buffer = buffer.split("\n\n", 1)
                for line in event_str.split("\n"):
                    if line.startswith("data: "):
                        try:
                            event = json.loads(line[6:])
                            if event.get("type") == "code_chunk":
                                full_code += event["content"]
                            if event.get("type") == "code_complete":
                                full_code = event["content"]
                            yield event
                        except json.JSONDecodeError:
                            pass

        yield {"type": "done", "content": "", "full_code": full_code}


def execute_code_stream(code: str,
                        url="http://127.0.0.1:" + str(config_data["server_port"]), session_id: str = ""):
    with httpx.stream("POST", url + "/api/exe-code/stream/",
                      json={"code": code, "session_id": session_id},
                      timeout=300.0) as response:
        if response.status_code != 200:
            yield {"type": "error", "content": f"HTTP {response.status_code}"}
            return

        full_ans = ""
        buffer = ""
        for chunk in response.iter_text():
            buffer += chunk
            while "\n\n" in buffer:
                event_str, buffer = buffer.split("\n\n", 1)
                for line in event_str.split("\n"):
                    if line.startswith("data: "):
                        try:
                            event = json.loads(line[6:])
                            if event.get("type") == "chunk":
                                full_ans += event["content"]
                            yield event
                        except json.JSONDecodeError:
                            pass

        yield {"type": "done", "content": "", "full_ans": full_ans}


def generate_and_execute_stream(question: str, tables: Optional[List[str]] = None,
                                 selected_fields: Optional[dict] = None,
                                 selected_functions: Optional[List[str]] = None,
                                 url="http://127.0.0.1:" + str(config_data["server_port"]), session_id: str = ""):
    payload = {"question": question, "tables": tables, "session_id": session_id}
    if selected_fields:
        payload["selected_fields"] = selected_fields
    if selected_functions:
        payload["selected_functions"] = selected_functions
    with httpx.stream("POST", url + "/api/generate-and-execute/stream/",
                      json=payload,
                      timeout=300.0) as response:
        if response.status_code != 200:
            yield {"type": "error", "content": f"HTTP {response.status_code}"}
            return
        buffer = ""
        for chunk in response.iter_text():
            buffer += chunk
            while "\n\n" in buffer:
                event_str, buffer = buffer.split("\n\n", 1)
                for line in event_str.split("\n"):
                    if line.startswith("data: "):
                        try:
                            event = json.loads(line[6:])
                            yield event
                        except json.JSONDecodeError:
                            pass


def step_chat_api_stream(question: str, tables: Optional[List[str]] = None,
                         selected_fields: Optional[dict] = None,
                         url="http://127.0.0.1:" + str(config_data["server_port"]), session_id: str = ""):
    payload = {"question": question, "session_id": session_id}
    if tables:
        payload["tables"] = tables
    if selected_fields:
        payload["selected_fields"] = selected_fields
    with httpx.stream("POST", url + "/api/step-chat/stream/",
                      json=payload,
                      timeout=300.0) as response:
        if response.status_code != 200:
            yield {"type": "error", "content": f"HTTP {response.status_code}"}
            return

        full_ans = ""
        buffer = ""
        for chunk in response.iter_text():
            buffer += chunk
            while "\n\n" in buffer:
                event_str, buffer = buffer.split("\n\n", 1)
                for line in event_str.split("\n"):
                    if line.startswith("data: "):
                        try:
                            event = json.loads(line[6:])
                            if event.get("type") == "chunk":
                                full_ans += event["content"]
                            yield event
                        except json.JSONDecodeError:
                            pass

        yield {"type": "done", "content": "", "full_ans": full_ans}


def plain_chat_api_stream(question: str, tables: Optional[List[str]] = None,
                          selected_fields: Optional[dict] = None,
                          url="http://127.0.0.1:" + str(config_data["server_port"]), session_id: str = ""):
    payload = {"question": question, "session_id": session_id}
    if tables:
        payload["tables"] = tables
    if selected_fields:
        payload["selected_fields"] = selected_fields
    with httpx.stream("POST", url + "/api/plain-chat/stream/",
                      json=payload,
                      timeout=300.0) as response:
        if response.status_code != 200:
            yield {"type": "error", "content": f"HTTP {response.status_code}"}
            return

        full_ans = ""
        buffer = ""
        for chunk in response.iter_text():
            buffer += chunk
            while "\n\n" in buffer:
                event_str, buffer = buffer.split("\n\n", 1)
                for line in event_str.split("\n"):
                    if line.startswith("data: "):
                        try:
                            event = json.loads(line[6:])
                            if event.get("type") == "chunk":
                                full_ans += event["content"]
                            yield event
                        except json.JSONDecodeError:
                            pass

        yield {"type": "done", "content": "", "full_ans": full_ans}


def filter_db_fields_stream(question: str, tables: Optional[List[str]] = None,
                            url="http://127.0.0.1:" + str(config_data["server_port"]), session_id: str = ""):
    payload = {"question": question, "session_id": session_id}
    if tables:
        payload["tables"] = tables
    with httpx.stream("POST", url + "/api/filter-db-fields/stream/",
                      json=payload,
                      timeout=300.0) as response:
        if response.status_code != 200:
            yield {"type": "error", "content": f"HTTP {response.status_code}"}
            return

        full_content = ""
        buffer = ""
        for chunk in response.iter_text():
            buffer += chunk
            while "\n\n" in buffer:
                event_str, buffer = buffer.split("\n\n", 1)
                for line in event_str.split("\n"):
                    if line.startswith("data: "):
                        try:
                            event = json.loads(line[6:])
                            if event.get("type") == "chunk":
                                full_content += event["content"]
                            yield event
                        except json.JSONDecodeError:
                            pass

        yield {"type": "done", "content": full_content}


def filter_functions_stream(question: str,
                            url="http://127.0.0.1:" + str(config_data["server_port"]), session_id: str = ""):
    with httpx.stream("POST", url + "/api/filter-functions/stream/",
                      json={"question": question, "session_id": session_id},
                      timeout=300.0) as response:
        if response.status_code != 200:
            yield {"type": "error", "content": f"HTTP {response.status_code}"}
            return

        full_content = ""
        buffer = ""
        for chunk in response.iter_text():
            buffer += chunk
            while "\n\n" in buffer:
                event_str, buffer = buffer.split("\n\n", 1)
                for line in event_str.split("\n"):
                    if line.startswith("data: "):
                        try:
                            event = json.loads(line[6:])
                            if event.get("type") == "chunk":
                                full_content += event["content"]
                            yield event
                        except json.JSONDecodeError:
                            pass


def upload_csv_api(file_content, table_name="uploaded_data"):
    url = f"http://127.0.0.1:{config_data['server_port']}/upload-csv/"
    files = {
        'file': ('data.csv', file_content, 'text/csv')
    }
    data = {
        'table_name': table_name
    }
    with httpx.Client(timeout=60.0) as client:
        try:
            response = client.post(url, files=files, data=data)
            if response.status_code == 200:
                return response.json()
            else:
                return {"error": f"{response.status_code}", "details": response.text}
        except httpx.RequestError as e:
            return {"error": f"{str(e)}"}


def upload_doc_api(file_content, filename, table_name="uploaded_data"):
    url = f"http://127.0.0.1:{config_data['server_port']}/upload-txt/"
    files = {
        'file': (filename, file_content, 'application/octet-stream')
    }
    data = {
        'table_name': table_name
    }
    with httpx.Client(timeout=60.0) as client:
        try:
            response = client.post(url, files=files, data=data)
            if response.status_code == 200:
                return response.json()
            else:
                return {"error": f"{response.status_code}", "details": response.text}
        except httpx.RequestError as e:
            return {"error": f"{str(e)}"}


def download_image(url):
    """Download web image"""
    try:
        response = requests.get(url, timeout=30)
        if response.status_code == 200:
            return io.BytesIO(response.content)
    except:
        pass
    return None


def markdown_to_word(doc, markdown_text):
    """Convert Markdown to Word document content"""
    html = markdown.markdown(markdown_text, extensions=['extra', 'tables'])
    soup = BeautifulSoup(html, 'html.parser')

    for element in soup.children:
        if element.name == 'h1':
            doc.add_heading(element.get_text(), level=1)
        elif element.name == 'h2':
            doc.add_heading(element.get_text(), level=2)
        elif element.name == 'h3':
            doc.add_heading(element.get_text(), level=3)
        elif element.name == 'h4':
            doc.add_heading(element.get_text(), level=4)
        elif element.name == 'p':
            p = doc.add_paragraph()
            for child in element.children:
                if child.name == 'strong' or child.name == 'b':
                    run = p.add_run(child.get_text())
                    run.bold = True
                elif child.name == 'em' or child.name == 'i':
                    run = p.add_run(child.get_text())
                    run.italic = True
                elif child.name == 'a':
                    run = p.add_run(child.get_text())
                    run.underline = True
                elif child.name == 'img':
                    img_url = child.get('src')
                    if img_url:
                        img_data = download_image(img_url)
                        if img_data:
                            try:
                                doc.add_picture(img_data, width=Inches(5))
                            except:
                                p.add_run(f"[Image: {img_url}]")
                        else:
                            p.add_run(f"[Image load failed: {img_url}]")
                elif child.string:
                    p.add_run(child.string)
                elif child.name is None:
                    if child.strip():
                        p.add_run(child.strip())
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                doc.add_paragraph(li.get_text(), style='List Bullet')
        elif element.name == 'ol':
            for idx, li in enumerate(element.find_all('li', recursive=False), 1):
                doc.add_paragraph(f"{idx}. {li.get_text()}", style='List Number')
        elif element.name == 'table':
            rows = element.find_all('tr')
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(['td', 'th'])))
                table.style = 'Table Grid'
                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell.get_text().strip()
        elif element.name == 'blockquote':
            p = doc.add_paragraph()
            p.add_run(element.get_text()).italic = True
            p.paragraph_format.left_indent = Inches(0.5)
        elif element.name == 'hr':
            doc.add_page_break()
        elif element.name == 'br':
            doc.add_paragraph()
        elif element.string and element.string.strip():
            doc.add_paragraph(element.string.strip())


def export_full_to_word(conversation_history):
    """Export full conversation history to Word document"""
    doc = Document()

    title = doc.add_heading('Data-Copilot Conversation Export (Full)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    from datetime import datetime
    doc.add_paragraph(f"Export Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()

    for entry in conversation_history:
        if isinstance(entry, dict):
            entry_type = entry.get("type", "")
            content = entry.get("content", "")
            if entry_type == "question":
                doc.add_heading('Question:', level=2)
                doc.add_paragraph(content)
            elif entry_type == "think":
                doc.add_heading('Plan:', level=2)
                markdown_to_word(doc, content)
            elif entry_type == "act":
                action = entry.get("action", "")
                code = entry.get("code", "")
                result = entry.get("result", "")
                error = entry.get("error", "")
                text = entry.get("text", "")
                solved_ans = entry.get("solved_ans", "")
                if code:
                    doc.add_heading('Generated Code:', level=2)
                    code_para = doc.add_paragraph()
                    code_run = code_para.add_run(code)
                    code_run.font.name = 'Courier New'
                    code_run.font.size = Pt(10)
                if result:
                    doc.add_heading('Execution Result:', level=2)
                    markdown_to_word(doc, result)
                if error:
                    doc.add_heading('Error:', level=2)
                    doc.add_paragraph(error)
                if text:
                    doc.add_heading(f'Output ({action}):', level=2)
                    markdown_to_word(doc, text)
                if solved_ans:
                    doc.add_heading('Solved Answer:', level=2)
                    markdown_to_word(doc, solved_ans)
            elif entry_type == "observe":
                doc.add_heading('Review:', level=2)
                markdown_to_word(doc, content)
            else:
                doc.add_paragraph(json.dumps(entry, ensure_ascii=False, indent=2))
        else:
            doc.add_paragraph(str(entry))

        doc.add_paragraph('_' * 50)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream


def export_essentials_to_word(conversation_history):
    """Export only answers and first question"""
    doc = Document()

    title = doc.add_heading('Data-Copilot Export', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    from datetime import datetime
    doc.add_paragraph(f"Export Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()

    first_question = None
    answers = []

    for entry in conversation_history:
        if isinstance(entry, dict):
            if entry.get("type") == "question" and first_question is None:
                first_question = entry.get("content", "")
            elif entry.get("type") == "act":
                text = entry.get("text", "")
                result = entry.get("result", "")
                solved_ans = entry.get("solved_ans", "")
                if text:
                    answers.append(text)
                if result:
                    answers.append(result)
                if solved_ans:
                    answers.append(solved_ans)
        else:
            answers.append(str(entry))

    if first_question:
        doc.add_paragraph(first_question)
        doc.add_paragraph()

    if answers:
        for answer in answers:
            markdown_to_word(doc, answer)
            doc.add_paragraph()

    if not first_question and not answers:
        doc.add_paragraph("No essential content found to export.")

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream